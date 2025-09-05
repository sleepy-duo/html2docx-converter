import logging
from docx import Document
from copy import deepcopy
from htmldocx import HtmlToDocx
import html

logging.basicConfig(level=logging.INFO)

def _insert_html_at_paragraph(para, html_text):
    """
    Parses HTML and inserts it at the location of a paragraph.
    Relies on the improved html2docx library to handle all HTML elements.
    """
    if not html_text:
        para.text = ""
        return

    try:
        parent = para._element.getparent()
        idx = parent.index(para._element)

        # Use a temporary document to parse the HTML
        temp_doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html_text, temp_doc)

        # Insert the parsed elements into the main document
        for element in temp_doc._body._element:
            parent.insert(idx, deepcopy(element))
            idx += 1

        # Remove the original placeholder paragraph
        parent.remove(para._element)

    except Exception as e:
        logging.error(f"Failed to insert HTML: {e}. HTML content: '{html.escape(html_text[:200])}...'")
        p = Document().add_paragraph("[Content generation failed, HTML parsing error]")
        parent.insert(idx, p._element)

if __name__ == "__main__":
    # Test HTML examples (same as before)
    test1_html = '''
    <h1>上海银行股份有限公司上下游情况分析</h1>\n\n<ul>\n  <li><strong>上游供应商情况</strong></li>\n  <table border=\"1\" class=\"dataframe\">\n    <thead>\n      <tr style=\"text-align: right;\">\n        <th>项目</th>\n        <th>客户名称</th>\n        <th>合作年限</th>\n        <th>采购商品类别</th>\n        <th>金额（万元）</th>\n        <th>结算方式</th>\n        <th>账期</th>\n        <th>业务量占比</th>\n      </tr>\n    </thead>\n    <tbody>\n      <tr>\n        <td>技术服务供应商</td>\n        <td>未公开</td>\n        <td>未明确</td>\n        <td>机房检测服务、IT基础设施维保</td>\n        <td>未明确</td>\n        <td>电汇/信用支付</td>\n        <td>项目周期1-3年</td>\n        <td>未明确</td>\n      </tr>\n      <tr>\n        <td>工程建设与装修供应商</td>\n        <td>未公开</td>\n        <td>未明确</td>\n        <td>园区楼层装修、物业外包服务</td>\n        <td>2472.9（康桥园区装修）</td>\n        <td>电汇/信用支付</td>\n        <td>按项目进度分期支付</td>\n        <td>未明确</td>\n      </tr>\n      <tr>\n        <td>设备与物资供应商</td>\n        <td>深圳齐心集团</td>\n        <td>未明确</td>\n        <td>营销物品采购</td>\n        <td>20.02（2025年项目）</td>\n        <td>电汇/信用支付</td>\n        <td>未明确</td>\n        <td>未明确</td>\n      </tr>\n    </tbody>\n  </table>\n  <p><strong>信息缺口说明</strong>：技术服务类供应商名称及合作年限需通过招标平台进一步确认；工程建设类供应商名称未公开，仅披露单笔金额；设备与物资采购中仅提及深圳齐心集团部分产品参与，整体供应链覆盖范围不完整。</p >\n</ul>\n\n<ul>\n  <li><strong>下游客户情况</strong></li>\n  <table border=\"1\" class=\"dataframe\">\n    <thead>\n      <tr style=\"text-align: right;\">\n        <th>项目</th>\n        <th>客户名称</th>\n        <th>合作年限</th>\n        <th>销售商品类别</th>\n        <th>金额（亿元）</th>\n        <th>结算方式</th>\n        <th>账期</th>\n        <th>业务量占比</th>\n      </tr>\n    </thead>\n    <tbody>\n      <tr>\n        <td>企业金融客户</td>\n        <td>泰禾集团</td>\n        <td>未明确</td>\n        <td>战略授信</td>\n        <td>200（2023年协议）</td>\n        <td>银行汇票/商业汇票</td>\n        <td>按合同约定（如按月付息）</td>\n        <td>约60%（行业推算）</td>\n      </tr>\n      <tr>\n        <td>个人金融客户</td>\n        <td>个人用户</td>\n        <td>未明确</td>\n        <td>财富管理、消费信贷</td>\n        <td>未明确</td>\n        <td>实时转账/按月结算</td>\n        <td>未明确</td>\n        <td>约40%（行业推算）</td>\n      </tr>\n      <tr>\n        <td>机构及政府客户</td>\n        <td>上海市政府相关机构</td>\n        <td>未明确</td>\n        <td>轨道交通、环保项目融资</td>\n        <td>未明确</td>\n        <td>电汇/跨境清算</td>\n        <td>受国际惯例及贸易条款影响</td>\n        <td>未明确</td>\n      </tr>\n      <tr>\n        <td>战略合作企业</td>\n        <td>腾讯、阿里、汇丰银行</td>\n        <td>2018年起（腾讯）/超10年（汇丰）</td>\n        <td>金融科技场景、跨境清算</td>\n        <td>未明确</td>\n        <td>电汇/信用支付</td>\n        <td>未明确</td>\n        <td>未明确</td>\n      </tr>\n    </tbody>\n  </table>\n  <p><strong>风险提示</strong>：2023年因13项违规被罚690万元，反映内控压力；房企客户（如泰禾集团）债务风险可能传导至银行。</p >\n  <p><strong>建议</strong>：需通过年报“关联交易”章节及监管处罚案例，进一步分析地产风险敞口及客户信息共享合规性。</p >\n</ul>\n\n<ul>\n  <li><strong>结算方式与账期</strong></li>\n  <p><strong>银行作为服务提供方</strong>：向客户提供汇票、本票、支票、电汇等金融工具服务，企业客户费用结算多采用实时转账或按月结算，贷款还款按合同约定（如按月付息、到期还本），跨境结算周期受国际惯例及贸易条款影响。</p >\n  <p><strong>银行自身采购结算</strong>：通过公开招标平台确定供应商后，采用电汇或信用支付。技术服务类可能按年度维保周期结算，工程建设类按项目进度分期支付（如康桥园区装修项目分阶段付款）。</p >\n  <p><strong>信息缺口</strong>：具体账期政策需联系银行官方客服或查阅年报“支付结算业务”章节确认。</p >\n</ul>
    '''
    test1_html = html.unescape(test1_html)
    
    test2_html = '''
    <h1>外部融资情况</h1>\n<ul>\n    <li><strong>2025年融资记录</strong>：根据公开信息（截至2025年9月5日），未发现宁波盈瑞聚合科技有限公司在2025年存在股权融资、债权融资或银行贷款等外部融资行为。所谓“2025年外部融资增资协议书”仅为通用合同模板，无实际交易佐证。</li>\n    <li><strong>2024年融资信息</strong>：2024年度报告显示公司注册资本与实缴资本均为2亿元人民币，但未披露当年是否通过外部融资补充资金。公开渠道未查询到其获得风险投资、银行授信或其他债务融资的明确记录。</li>\n    <li><strong>政策性融资适配性</strong>：公司持有高新技术企业、专精特新“小巨人”资质，且拥有24项专利及多项行政许可，符合宁波市“甬质贷”等政策性融资工具的支持条件。但截至2025年9月5日，无证据表明其已参与该计划或获得相关授信。</li>\n    <li><strong>潜在风险提示</strong>：公司存在1条司法案件及4条开庭公告，可能对融资信用评估产生影响，需进一步核实案件性质及处理进展。</li>\n</ul>\n\n<table border=\"1\" class=\"dataframe\">\n    <thead>\n        <tr style=\"text-align: right;\">\n            <th>融资类型</th>\n            <th>公开信息状态</th>\n            <th>备注</th>\n        </tr>\n    </thead>\n    <tbody>\n        <tr>\n            <td>股权融资</td>\n            <td>未披露</td>\n            <td>2025年无公开记录</td>\n        </tr>\n        <tr>\n            <td>债权融资</td>\n            <td>未披露</td>\n            <td>2025年无公开记录</td>\n        </tr>\n        <tr>\n            <td>银行贷款</td>\n            <td>未披露</td>\n            <td>需核实非公开渠道贷款情况</td>\n        </tr>\n        <tr>\n            <td>政策性融资（甬质贷）</td>\n            <td>未确认参与</td>\n            <td>资质匹配但无实际授信证据</td>\n        </tr>\n    </tbody>\n</table>\n\n<p><strong>建议行动</strong>：建议客户经理通过以下途径补充核实：<br>\n1. 查询国家企业信用信息公示系统，确认2024-2025年工商变更记录及注册资本调整情况；<br>\n2. 通过天眼查/企查查等平台追踪股东结构变动及未公开融资事件；<br>\n3. 联系宁波市金融监管局或“甬质贷”项目方，核实企业政策性融资参与度；<br>\n4. 调取司法案件详情，评估其对融资信用的实际影响。</p >
    '''
    test2_html = html.unescape(test2_html)
    
    # Create document
    # Make sure to have the 'intermediate.docx' file in the right path
    # Or change the path to an absolute path
    try:
        doc = Document("/Users/qixin/Downloads/due_deligence/data/reports/intermediate.docx")
    except Exception as e:
        print("Error opening intermediate.docx. Make sure the file exists.")
        print("You can create a blank docx file and name it intermediate.docx")
        print(e)
        exit()

    all_paras = [p for p in doc.paragraphs]
    # Create a copy of the list to iterate over, as we are modifying the document
    for para in all_paras[:]:        
        tag1 = f"__PLACEHOLDER_upstream_and_downstream__"
        tag2 = f"__PLACEHOLDER_group_overview__"
        if tag1 in para.text:
            # The new h2d.py should handle this complex html
            _insert_html_at_paragraph(para, test1_html)
        elif tag2 in para.text:
            # The new h2d.py should handle this complex html
            _insert_html_at_paragraph(para, test2_html)
        else:
            continue

    # Save file
    output_file = "test_output1.docx"
    doc.save(output_file)
    print(f"Test complete, {output_file} generated.")
