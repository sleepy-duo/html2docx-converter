"""
This module provides the HtmlToDocx class, the core of the html2docx library.

It uses BeautifulSoup to parse HTML and python-docx to create .docx files,
with a focus on handling complex structures like nested lists and merged table cells.
"""

import re, argparse
import io, os
import urllib.request
from urllib.parse import urlparse
from html.parser import HTMLParser

import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from bs4 import BeautifulSoup, NavigableString

# Constants for styling
INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5
DEFAULT_TABLE_STYLE = 'TableGrid'
DEFAULT_PARAGRAPH_STYLE = None


def get_filename_from_url(url):
    """Extracts the filename from a URL."""
    return os.path.basename(urlparse(url).path)


def is_url(url):
    """Checks if a string is a URL."""
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])


def fetch_image(url):
    """Fetches an image from a URL and returns it as a BytesIO object."""
    try:
        with urllib.request.urlopen(url) as response:
            return io.BytesIO(response.read())
    except urllib.error.URLError:
        return None


def delete_paragraph(paragraph):
    """Deletes a paragraph from a document."""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def to_roman(n):
    """Converts an integer to a Roman numeral."""
    if not 1 <= n < 4000:
        return str(n)
    roman_map = [
        (1000, 'm'), (900, 'cm'), (500, 'd'), (400, 'cd'),
        (100, 'c'), (90, 'xc'), (50, 'l'), (40, 'xl'),
        (10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'),
        (1, 'i')
    ]
    result = []
    for value, numeral in roman_map:
        while n >= value:
            result.append(numeral)
            n -= value
    return ''.join(result)


# Mappings from HTML tags to python-docx font styles
font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

# Mappings from HTML tags to font names
font_names = {
    'code': 'Courier New',
    'pre': 'Courier New',
}


class HtmlToDocx(HTMLParser):
    """
    Parses HTML and converts it to a .docx document.

    This class inherits from HTMLParser, but the parsing is primarily done using
    BeautifulSoup for more robust handling of malformed HTML.
    """

    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_style = DEFAULT_TABLE_STYLE
        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE

    def set_initial_attrs(self, document=None):
        """Set initial attributes for a new parsing session."""
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options['fix-html']
        self.document = self.doc
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.run = None

    def copy_settings_from(self, other):
        """Copy settings from another parser instance."""
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def parse_dict_string(self, string, separator=';'):
        """Parses a CSS-style string into a dictionary."""
        new_string = string.replace(" ",'').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def add_styles_to_paragraph(self, style):
        """Applies CSS styles to a paragraph."""
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9.]+', '', margin)
            margin = float(re.sub(r'[a-z]+', '', margin))
            if units == 'px':
                self.paragraph.paragraph_format.left_indent = Inches(min(margin / 10 * INDENT, MAX_INDENT))

    def add_styles_to_run(self, style):
        """Applies CSS styles to a run."""
        if 'color' in style:
            if 'rgb' in style['color']:
                color = re.sub(r'[a-z()]+', '', style['color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['color']:
                color = style['color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
            self.run.font.color.rgb = RGBColor(*colors)
            
        if 'background-color' in style:
            if 'rgb' in style['background-color']:
                color = re.sub(r'[a-z()]+', '', style['background-color'])
                colors = [int(x) for x in color.split(',')]
            elif '#' in style['background-color']:
                color = style['background-color'].lstrip('#')
                colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = [0, 0, 0]
            self.run.font.highlight_color = WD_COLOR.GRAY_25

    def apply_paragraph_style(self, style=None):
        """Applies a named paragraph style to the current paragraph."""
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def add_html_to_document(self, html, document):
        """
        Parses the given HTML and adds the content to the given document.

        :param str html: The HTML content to parse.
        :param docx.document.Document document: The document to add the content to.
        """
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a string.')
        if not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a docx Document or Cell object.')
        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        """
        Parses the given HTML and adds the content to the given table cell.

        This is used for parsing content within table cells, including nested tables.
        :param str html: The HTML content to parse.
        :param docx.table._Cell cell: The cell to add the content to.
        """
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a docx Cell object.')
        # Clear the cell's content before adding new content
        cell.text = ''
        unwanted_paragraph = cell.paragraphs[0]
        delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # Cells must end with a paragraph or will get message about corrupt file
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    def parse_html_file(self, filename_html, filename_docx=None):
        """
        Parses an HTML file and saves the content to a .docx file.

        :param str filename_html: The path to the HTML file.
        :param str filename_docx: The path to the output .docx file. Optional.
        """
        with open(filename_html, 'r', encoding='utf-8') as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = f'{os.path.splitext(filename)[0]}.docx'
        self.doc.save(filename_docx)
    
    def parse_html_string(self, html):
        """
        Parses an HTML string and returns a new docx.Document object.

        :param str html: The HTML content to parse.
        :return: A new docx.Document object.
        :rtype: docx.document.Document
        """
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc

    def run_process(self, html):
        """Orchestrates the parsing process."""
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            body = self.soup.find('body')
            if body:
                elements = body.contents
            else:
                elements = self.soup.contents
            
            self._parse_elements(elements, self.doc)
        else:
            # Fallback to original feed parser if bs4 is not available
            self.feed(html)

    def _parse_elements(self, elements, parent):
        """
        Recursively parses a list of BeautifulSoup elements and adds them to the parent.

        The parent can be a Document, a Cell, or a Paragraph object.
        """
        for element in elements:
            if isinstance(element, NavigableString):
                if element.strip():
                    if isinstance(parent, docx.text.paragraph.Paragraph):
                        parent.add_run(element.strip())
                    else:
                        if parent.paragraphs:
                            p = parent.paragraphs[-1]
                        else:
                            p = parent.add_paragraph()
                        p.add_run(element.strip())
            elif element.name:
                if element.name == 'table':
                    self.handle_table(element, parent)
                elif element.name in ['ul', 'ol']:
                    self.handle_list(element, parent)
                elif element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.handle_paragraph(element, parent)
                elif element.name == 'blockquote':
                    self.handle_blockquote(element, parent)
                elif element.name == 'pre':
                    self.handle_pre(element, parent)
                elif element.name == 'hr':
                    self.handle_hr(parent)
                elif element.name == 'br':
                    if isinstance(parent, docx.document.Document):
                        p = parent.add_paragraph()
                    else:
                        p = parent.paragraphs[-1] if parent.paragraphs else parent.add_paragraph()
                    if p.runs:
                        p.runs[-1].add_break()
                elif element.name == 'img':
                    self.handle_img(element.attrs, parent)
                elif element.name in font_styles or element.name in font_names:
                    self.handle_inline_style(element, parent)
                elif element.name == 'a':
                    self.handle_link(element.attrs.get('href', ''), element.get_text(strip=True), parent)
                else:
                    # For unknown tags, recursively parse their children
                    self._parse_elements(element.contents, parent)

    def handle_table(self, table_tag, parent):
        """
        Handles `<table>` tags, including merged cells and header rows.
        A new HtmlToDocx instance is created for each cell to handle nested content in isolation.
        """
        rows = []
        # Find all `tr` elements, whether they are direct children of the table
        # or nested inside `thead`, `tbody`, or `tfoot`.
        for child in table_tag.children:
            if isinstance(child, NavigableString):
                continue
            if child.name == 'tr':
                rows.append(child)
            elif child.name in ['thead', 'tbody', 'tfoot']:
                rows.extend(child.find_all('tr', recursive=False))

        if not rows:
            return

        num_rows = len(rows)
        max_cols = 0
        for r in rows:
            current_row_cols = 0
            for cell in r.find_all(['td', 'th'], recursive=False):
                current_row_cols += int(cell.get('colspan', 1))
            if current_row_cols > max_cols:
                max_cols = current_row_cols
        
        if max_cols == 0:
            return

        table = self.doc.add_table(rows=num_rows, cols=max_cols)
        table.style = self.table_style
        grid_map = [[None for _ in range(max_cols)] for _ in range(num_rows)]

        for i, row in enumerate(rows):
            cols = row.find_all(["td", "th"], recursive=False)
            is_header_row = (row.parent.name == 'thead')
            j = 0
            for cell in cols:
                while j < max_cols and grid_map[i][j] is not None:
                    j += 1
                
                r_span = int(cell.get("rowspan", 1))
                c_span = int(cell.get("colspan", 1))

                if i + r_span > num_rows: r_span = num_rows - i
                if j + c_span > max_cols: c_span = max_cols - j

                docx_cell = table.cell(i, j)
                
                # Create a new parser for each cell to handle nested content correctly.
                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell.decode_contents(), docx_cell)

                # Style header cells
                if is_header_row or cell.name == 'th':
                    for p in docx_cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                    tc = docx_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'D9D9D9') # Light grey
                    tcPr.append(shd)

                if r_span > 1 or c_span > 1:
                    end_cell = table.cell(i + r_span - 1, j + c_span - 1)
                    docx_cell.merge(end_cell)

                for r in range(i, i + r_span):
                    for c in range(j, j + c_span):
                        if r < num_rows and c < max_cols:
                            grid_map[r][c] = True
    
    def handle_list(self, list_tag, parent, level=0, start=1):
        """
        Handles `<ul>` and `<ol>` tags, including nested lists and malformed HTML.
        Different bullet and numbering styles are used for different list levels.
        """
        item_counter = start
        ul_bullets = ['•', 'o', '-']

        for child in list_tag.children:
            if isinstance(child, NavigableString):
                continue
            
            if child.name == 'li':
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(min((level + 1) * LIST_INDENT, MAX_INDENT))
                
                if list_tag.name == 'ol':
                    if level == 0:
                        num_str = f'{item_counter}.'
                    elif level == 1:
                        num_str = f'{chr(ord("a") + item_counter - 1)}.'
                    else:
                        num_str = f'{to_roman(item_counter)}.'
                    run = p.add_run(f'{num_str}\t')
                    item_counter += 1
                else:
                    bullet = ul_bullets[level % len(ul_bullets)]
                    run = p.add_run(f'{bullet}\t')
                
                self._parse_list_item_content(child, p, level)

            elif child.name in ['ul', 'ol']:
                # Handles malformed lists where a list is a direct child of another list.
                self.handle_list(child, parent, level + 1)

    def _parse_list_item_content(self, item, p, level):
        """Parses the content of an `<li>` element."""
        for child in item.children:
            if isinstance(child, NavigableString):
                if child.strip():
                    p.add_run(child.strip())
            elif child.name in ['ul', 'ol']:
                self.handle_list(child, self.doc, level + 1)
            elif child.name in font_styles or child.name in font_names:
                self.handle_inline_style(child, p)
            elif child.name == 'a':
                self.handle_link(child.attrs.get('href', ''), child.get_text(strip=True), p)
            elif child.name == 'br':
                p.add_run().add_break()
            else:
                # For unknown tags, just add their text content.
                p.add_run(child.get_text(strip=True))

    def handle_paragraph(self, p_tag, parent):
        """Handles `<p>` and heading tags."""
        p = self.doc.add_paragraph()
        if p_tag.name.startswith('h'):
            try:
                level = int(p_tag.name[1])
                p.style = f'Heading {level}'
                p.paragraph_format.space_after = Pt(6)
            except (ValueError, KeyError):
                self.apply_paragraph_style()
        else:
            self.apply_paragraph_style()
            p.paragraph_format.space_after = Pt(2)

        if 'style' in p_tag.attrs:
            style = self.parse_dict_string(p_tag['style'])
            self.paragraph = p
            self.add_styles_to_paragraph(style)
        
        self._parse_elements(p_tag.contents, p)

    def handle_blockquote(self, bq_tag, parent):
        """Handles `<blockquote>` tags with indentation and a background color."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0') # Light grey
        pPr.append(shd)

        self._parse_elements(bq_tag.contents, p)

    def handle_pre(self, pre_tag, parent):
        """Handles `<pre>` tags, preserving whitespace and using a monospace font."""
        p = self.doc.add_paragraph()
        p.style = 'No Spacing'
        
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5') # Very light grey
        pPr.append(shd)

        text = pre_tag.get_text()
        run = p.add_run(text)
        run.font.name = 'Courier New'

    def handle_hr(self, parent):
        """Handles `<hr>` tags by adding a paragraph with a bottom border."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def handle_inline_style(self, element, parent):
        """Handles inline style tags like `<b>`, `<i>`, `<code>`, etc."""
        if isinstance(parent, docx.text.paragraph.Paragraph):
            p = parent
        elif parent.paragraphs:
            p = parent.paragraphs[-1]
        else:
            p = parent.add_paragraph()
        
        run = p.add_run(element.get_text(strip=True))
        if element.name in font_styles:
            setattr(run.font, font_styles[element.name], True)
        if element.name in font_names:
            run.font.name = font_names[element.name]

    def handle_link(self, href, text, parent):
        """Handles `<a>` tags."""
        if not href:
            if isinstance(parent, docx.text.paragraph.Paragraph):
                p = parent
            elif parent.paragraphs:
                p = parent.paragraphs[-1]
            else:
                p = parent.add_paragraph()
            p.add_run(text)
            return

        if isinstance(parent, docx.document.Document):
            p = parent.add_paragraph()
        elif isinstance(parent, docx.table._Cell):
            p = parent.add_paragraph()
        else: # It should be a paragraph
            p = parent

        is_external = href.startswith('http')
        rel_id = p.part.relate_to(
            href,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True
        )

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), rel_id)

        subrun = p.add_run()
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), "0000EE")
        rPr.append(c)

        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

        subrun._r.append(rPr)
        subrun._r.text = text
        hyperlink.append(subrun._r)
        p._p.append(hyperlink)

    def handle_img(self, attrs, parent):
        """Handles `<img>` tags."""
        if not self.include_images:
            return
        src = attrs.get('src')
        if not src:
            return
            
        src_is_url = is_url(src)
        image = None
        if src_is_url:
            image = fetch_image(src)
        elif os.path.exists(src):
            image = src

        if image:
            try:
                if hasattr(parent, 'add_picture'):
                    parent.add_picture(image)
                elif hasattr(parent, 'add_paragraph'):
                    p = parent.add_paragraph() if not isinstance(parent, docx.text.paragraph.Paragraph) else parent
                    r = p.add_run()
                    r.add_picture(image)

            except Exception:
                image = None
        
        if not image:
            text = f"<image: {src}>" if src_is_url else f"<image: {get_filename_from_url(src)}>"
            if hasattr(parent, 'add_paragraph'):
                 parent.add_paragraph(text)


if __name__=="__main__":
    arg_parser = argparse.ArgumentParser(description='Convert .html file into .docx file with formatting')
    arg_parser.add_argument('filename_html', help='The .html file to be parsed')
    arg_parser.add_argument(
        'filename_docx', 
        nargs='?', 
        help='The name of the .docx file to be saved. Default new_docx_file_[filename_html]', 
        default=None
    )
    arg_parser.add_argument('--bs', action='store_true', 
        help='Attempt to fix html before parsing. Requires bs4. Default True')

    args = vars(arg_parser.parse_args())
    file_html = args.pop('filename_html')
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(file_html, **args)
