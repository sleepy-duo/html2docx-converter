import logging
from docx import Document
from copy import deepcopy
from htmldocx import HtmlToDocx
import html

logging.basicConfig(level=logging.INFO)

def _insert_html_at_paragraph(para, html_text):
    """
    Parses HTML and inserts it at the location of a paragraph.
    Relies on the improved html2docx library to handle all HTML elements.
    """
    if not html_text:
        para.text = ""
        return

    try:
        parent = para._element.getparent()
        idx = parent.index(para._element)

        # Use a temporary document to parse the HTML
        temp_doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html_text, temp_doc)

        # Insert the parsed elements into the main document
        for element in temp_doc._body._element:
            parent.insert(idx, deepcopy(element))
            idx += 1

        # Remove the original placeholder paragraph
        parent.remove(para._element)

    except Exception as e:
        logging.error(f"Failed to insert HTML: {e}. HTML content: '{html.escape(html_text[:200])}...'")
        p = Document().add_paragraph("[Content generation failed, HTML parsing error]")
        parent.insert(idx, p._element)

if __name__ == "__main__":
    # Test HTML examples (same as before)
    test1_html = '''
    <h1>浦诺菲新材料有限公司访谈提纲</h1>

<ul>
  <li><strong>一、公司治理与股东结构</strong></li>
  <ul>
    <li>请说明集团整体法律架构及实际控制人认定依据</li>
    <li>补充披露上海晟世优新材料科技有限公司、宁波膜杰材料科技有限公司等主要股东的实缴资本证明</li>
    <li>请提供宁波激智科技股份有限公司的技术合作细节及供应链协同机制</li>
    <li>核实宁波浦诺玉成合伙企业2029年认缴出资的可行性及资金来源</li>
  </ul>

  <li><strong>二、财务与融资需求</strong></li>
  <ul>
    <li>请提供2022-2025年完整财务报表（含附注）及审计报告</li>
    <li>说明2021年1.21亿元应收账款的形成原因及后续改善措施</li>
    <li>披露近三年研发费用总额及资本化比例，验证技术投入产出比</li>
    <li>请提供2024年苏州基地建设资金来源及资产负债率数据</li>
    <li>说明与上海银行"万企千亿行动"潜在合作的具体需求场景</li>
  </ul>

  <li><strong>三、业务与市场拓展</strong></li>
  <ul>
    <li>请提供2024年国内经销商（广汇、永达等）订单占比及信用评级</li>
    <li>说明国际主机厂（宝马、奥迪等）合作模式及订单稳定性</li>
    <li>披露前装市场（OEM）客户验证进展及渗透率数据</li>
    <li>请提供2025年线上渠道（天猫等）销售占比及获客成本结构</li>
    <li>说明应对纳米复合膜、固态电解质等技术替代的具体研发规划</li>
  </ul>

  <li><strong>四、供应链与生产管理</strong></li>
  <ul>
    <li>请提供前五大供应商名单及采购集中度数据</li>
    <li>说明高性能基膜国产化替代进展及供应链韧性建设方案</li>
    <li>披露PET原膜挤出、纳米涂布等核心技术的专利布局情况</li>
    <li>请提供宁波、苏州生产基地产能利用率及良品率数据</li>
  </ul>

  <li><strong>五、风险与合规管理</strong></li>
  <ul>
    <li>请说明上海子公司司法案件的具体情况及处理进展</li>
    <li>提供IATF16949车规质量认证、环保合规文件及危险化学品使用情况</li>
    <li>说明欧盟REACH标准等国际环保政策的应对方案</li>
    <li>请提供近三年重大合同履约情况及供应商管理机制</li>
  </ul>
</ul>

<table border="1">
  <tr>
    <th>核心信息缺口</th>
    <th>建议调查方向</th>
    <th>数据来源建议</th>
  </tr>
  <tr>
    <td>集团法律架构</td>
    <td>核实实际控制人及集团名称</td>
    <td>企业信用信息公示系统</td>
  </tr>
  <tr>
    <td>财务数据</td>
    <td>获取经审计的三表数据及附注</td>
    <td>企业年报/审计报告</td>
  </tr>
  <tr>
    <td>客户信用管理</td>
    <td>补充国内经销商信用评级</td>
    <td>客户征信报告</td>
  </tr>
  <tr>
    <td>技术替代应对</td>
    <td>验证技术路线迭代规划</td>
    <td>研发立项文件</td>
  </tr>
  <tr>
    <td>供应链风险</td>
    <td>评估基膜国产化替代进展</td>
    <td>采购合同台账</td>
  </tr>
</table>

<ul>
  <li><strong>六、战略与行业趋势</strong></li>
  <ul>
    <li>请说明在新能源车轻量化赛道的技术突破方向</li>
    <li>披露与主机厂数... [truncated]
    '''
    test1_html = html.unescape(test1_html)
    
    test2_html = '''
    <h1>三、上下游客户情况分析</h1>
<ul>
  <li><strong>上游供应商信息缺失</strong>：背景知识中未提及具体供应商名称、业务集中度及结算方式，建议补充供应链调查。</li>
  <li><strong>下游客户集中度较高</strong>：主要客户为半导体制造龙头企业，前五名客户收入占比62%-69%（2020-2023年数据）。</li>
</ul>

<table border="1" cellpadding="5">
  <caption>上下游合作情况汇总</caption>
  <thead>
    <tr>
      <th>项目</th>
      <th>客户名称</th>
      <th>采购/销售商品类别</th>
      <th>业务量占比</th>
      <th>结算方式</th>
      <th>账期</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td rowspan="5">销售环节（下游）</td>
      <td>中芯国际</td>
      <td rowspan="5">电子专用材料、化学品</td>
      <td>60%（2018年）</td>
      <td rowspan="5">无明确信息</td>
      <td rowspan="5">无明确信息</td>
    </tr>
    <tr>
      <td>台积电</td>
      <td>8%（2018年）</td>
    </tr>
    <tr>
      <td>长江存储</td>
      <td>8%（2018年）</td>
    </tr>
    <tr>
      <td>华润微电</td>
      <td>前五名合计占比62%-69%</td>
    </tr>
    <tr>
      <td>华虹宏力</td>
      <td>前五名合计占比62%-69%</td>
    </tr>
    <tr>
      <td>采购环节（上游）</td>
      <td colspan="5" style="text-align:center">背景知识未提供供应商及结算信息</td>
    </tr>
  </tbody>
</table>

<ul>
  <li><strong>结算方式说明</strong>：背景知识中未披露与上下游的具体结算条款，需核查企业购销合同。</li>
  <li><strong>风险提示</strong>：下游客户集中度高于60%，需重点关注大客户依赖风险及回款稳定性。</li>
</ul>

    '''
    test2_html = html.unescape(test2_html)
    
    # Create document
    # Make sure to have the 'intermediate.docx' file in the right path
    # Or change the path to an absolute path
    try:
        doc = Document("/Users/qixin/Downloads/due_deligence/data/reports/intermediate.docx")
    except Exception as e:
        print("Error opening intermediate.docx. Make sure the file exists.")
        print("You can create a blank docx file and name it intermediate.docx")
        print(e)
        exit()

    all_paras = [p for p in doc.paragraphs]
    # Create a copy of the list to iterate over, as we are modifying the document
    for para in all_paras[:]:        
        tag1 = f"__PLACEHOLDER_upstream_and_downstream__"
        tag2 = f"__PLACEHOLDER_group_overview__"
        if tag1 in para.text:
            # The new h2d.py should handle this complex html
            _insert_html_at_paragraph(para, test1_html)
        elif tag2 in para.text:
            # The new h2d.py should handle this complex html
            _insert_html_at_paragraph(para, test2_html)
        else:
            continue

    # Save file
    output_file = "test_output1.docx"
    doc.save(output_file)
    print(f"Test complete, {output_file} generated.")
